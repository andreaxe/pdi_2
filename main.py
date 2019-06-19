import os
import pandas as pd
from network import build_network
from sqlalchemy import create_engine
from dataset import generate_data

# =============================================
# Definições de base de dados
# =============================================
username = 'root'
password = ''
hostname = '127.0.0.1'
database = 'pdi_par'
# =============================================

# Salvar a informação sobre a rede (usado) para o relatório
save_network = True


def doc_table(data, table_title, path_file):
    """

    :param data:
    :param table_title:
    :param path_file:
    :return:
    """

    from docx import Document
    import pandas as pd
    document = Document()
    data = pd.DataFrame(data)  # My input data is in the 2D list form
    document.add_heading(table_title)
    table = document.add_table(rows=(data.shape[0]), cols=data.shape[1])  # First row are table headers!
    for i, column in enumerate(data):
        for row in range(data.shape[0]):
            table.cell(row, i).text = str(data[column][row])
    document.save(path_file)


if __name__ == '__main__':

    dim_arr  = []
    rdim_arr = []

    engine = create_engine("mysql+pymysql://{}:{}@{}/{}".format(username, password, hostname, database))
    con = engine.connect()

    # ========================================================================
    # this function generates a random dataset and re-populates the database
    generate_data(number_of_nodes=20, con=con)
    # ========================================================================

    sql = 'select * from dim;'
    results = con.execute(sql)
    result_set = results.fetchall()

    for row in result_set:
        dim_arr.append(dict(row))

    sql = 'select * from rdim;'
    results = con.execute(sql)
    result_set = results.fetchall()

    for row in result_set:
        rdim_arr.append(dict(row))

    print(rdim_arr)

    if save_network:

        dim  = pd.DataFrame(dim_arr)
        rdim = pd.DataFrame(rdim_arr)

        doc_table(dim,  'Tabela DIM',  os.path.join('documents', 'dim.docx'))
        doc_table(rdim, 'Tabela RDIM', os.path.join('documents', 'rdim.docx'))

    build_network(dim_arr, rdim_arr)





